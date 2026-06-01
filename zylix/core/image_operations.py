import sys
from pathlib import Path

if getattr(sys, 'frozen', False):
    sys.path.insert(0, str(Path(sys.executable).parent))

from pathlib import Path
from typing import List
import numpy as np

from PIL import Image

from .base import OperationBase, OperationResult, ImageOperationBase


class ImageToPdf(ImageOperationBase):

    @property
    def name(self) -> str:
        return "Imagen a PDF"

    @property
    def input_extensions(self) -> List[str]:
        return [".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"]

    @property
    def output_extension(self) -> str:
        return ".pdf"

    def execute(self, input_paths: List[Path], output_dir: Path, **kwargs) -> OperationResult:
        from datetime import datetime
        output_files = []
        mode = kwargs.get("mode", "separate")
        output_name = kwargs.get("output_name", "").strip()

        try:
            if mode == "joined" and output_name:
                images = []
                for img_path in input_paths:
                    img = Image.open(img_path).convert("RGB")
                    images.append(img)

                if images:
                    pdf_path = output_dir / f"{output_name}.pdf"
                    images[0].save(pdf_path, "PDF", resolution=kwargs.get("dpi", 300), save_all=True, append_images=images[1:])
                    output_files.append(pdf_path)
            else:
                timestamp = datetime.now().strftime("%H%M%S")
                for img_path in input_paths:
                    img = Image.open(img_path).convert("RGB")
                    pdf_path = output_dir / f"{img_path.stem}_{timestamp}.pdf"
                    img.save(pdf_path, "PDF", resolution=kwargs.get("dpi", 300))
                    output_files.append(pdf_path)

            return OperationResult(True, f"Creado {len(output_files)} PDF(s)", output_files)
        except Exception as e:
            return OperationResult(False, f"Error: {str(e)}")


class ConvertImageFormat(ImageOperationBase):

    @property
    def name(self) -> str:
        return "Cambiar Formato"

    @property
    def output_extension(self) -> str:
        return ".png"

    def execute(self, input_paths: List[Path], output_dir: Path, **kwargs) -> OperationResult:
        output_format = kwargs.get("format", "PNG").upper()
        output_files = []

        try:
            for img_path in input_paths:
                img = Image.open(img_path)
                if img.mode == "RGBA" and output_format in ["JPEG", "JPG"]:
                    img = img.convert("RGB")
                elif img.mode != "RGB" and output_format in ["JPEG", "JPG", "WEBP"]:
                    img = img.convert("RGB")

                output_path = output_dir / f"{img_path.stem}.{output_format.lower()}"
                save_format = "JPEG" if output_format == "JPG" else output_format
                img.save(output_path, format=save_format)
                output_files.append(output_path)

            return OperationResult(True, f"Convertido {len(output_files)} imagen(es)", output_files)
        except Exception as e:
            return OperationResult(False, f"Error: {str(e)}")


class OptimizeImage(ImageOperationBase):

    @property
    def name(self) -> str:
        return "Optimizar Imagen"

    @property
    def output_extension(self) -> str:
        return ".jpg"

    def execute(self, input_paths: List[Path], output_dir: Path, **kwargs) -> OperationResult:
        quality = kwargs.get("quality", 85)
        output_files = []

        try:
            for img_path in input_paths:
                img = Image.open(img_path)
                if img.mode == "RGBA":
                    rgb = Image.new("RGB", img.size, (255, 255, 255))
                    rgb.paste(img, mask=img.split()[3])
                    img = rgb

                output_path = output_dir / f"{img_path.stem}_optimized.jpg"
                img.save(output_path, "JPEG", quality=quality, optimize=True)
                output_files.append(output_path)

            return OperationResult(True, f"Optimizado {len(output_files)} imagen(es)", output_files)
        except Exception as e:
            return OperationResult(False, f"Error: {str(e)}")


class RotateImage(ImageOperationBase):

    @property
    def name(self) -> str:
        return "Rotar Imagen"

    @property
    def output_extension(self) -> str:
        return ".png"

    def execute(self, input_paths: List[Path], output_dir: Path, **kwargs) -> OperationResult:
        angle = kwargs.get("angle", 90)
        output_files = []

        try:
            for img_path in input_paths:
                img = Image.open(img_path)
                rotated = img.rotate(angle, expand=True, fillcolor=(255, 255, 255))

                output_path = output_dir / f"{img_path.stem}_rotated{angle}.png"
                rotated.save(output_path)
                output_files.append(output_path)

            return OperationResult(True, f"Rotado {len(output_files)} imagen(es)", output_files)
        except Exception as e:
            return OperationResult(False, f"Error: {str(e)}")


class RemoveBackground(ImageOperationBase):

    @property
    def name(self) -> str:
        return "Remover Fondo"

    @property
    def output_extension(self) -> str:
        return ".png"

    def execute(self, input_paths: List[Path], output_dir: Path, **kwargs) -> OperationResult:
        from skimage import io, color, morphology
        output_files = []

        try:
            for img_path in input_paths:
                img = io.imread(img_path)
                if len(img.shape) == 2:
                    img = color.gray2rgb(img)
                if img.shape[-1] == 4:
                    img = img[:,:,:3]

                gray = color.rgb2gray(img)
                threshold = np.mean(gray) + 0.1
                mask = (gray < threshold).astype(np.uint8) * 255

                mask = morphology.opening(mask, morphology.disk(3))
                mask = morphology.closing(mask, morphology.disk(5))

                result = np.zeros_like(img)
                result[mask == 255] = img[mask == 255]

                output_path = output_dir / f"{img_path.stem}_nobg.png"
                io.imsave(output_path, result)
                output_files.append(output_path)

            return OperationResult(True, f"Fondo removido de {len(output_files)} imagen(es)", output_files)
        except Exception as e:
            return OperationResult(False, f"Error: {str(e)}")


class ExtractTextFromImage(ImageOperationBase):

    @property
    def name(self) -> str:
        return "OCR de Imagen"

    @property
    def output_extension(self) -> str:
        return ".txt"

    def execute(self, input_paths: List[Path], output_dir: Path, **kwargs) -> OperationResult:
        from .ocr_engine import OCREngine

        engine = OCREngine()
        output_files = []

        try:
            for img_path in input_paths:
                text = engine.extract_text(img_path)

                output_path = output_dir / f"{img_path.stem}_ocr.txt"
                output_path.write_text(text, encoding="utf-8")
                output_files.append(output_path)

            return OperationResult(True, f"OCR realizado en {len(output_files)} imagen(es)", output_files)
        except Exception as e:
            return OperationResult(False, f"Error: {str(e)}")


class ImageToWord(ImageOperationBase):

    @property
    def name(self) -> str:
        return "Imagen a Word"

    @property
    def input_extensions(self) -> List[str]:
        return [".jpg", ".jpeg", ".png", ".bmp", ".gif", ".webp"]

    @property
    def output_extension(self) -> str:
        return ".docx"

    def execute(self, input_paths: List[Path], output_dir: Path, **kwargs) -> OperationResult:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        output_files = []
        add_blank_page = kwargs.get("add_blank_page", False)
        output_name = kwargs.get("output_name", "").strip()

        try:
            doc = Document()
            section = doc.sections[0]
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            max_width = section.page_width - section.left_margin - section.right_margin
            max_height = section.page_height - section.top_margin - section.bottom_margin

            if add_blank_page:
                doc.add_page_break()

            for i, img_path in enumerate(sorted(input_paths, key=lambda p: p.name.lower())):
                with Image.open(img_path) as img:
                    ancho_px, alto_px = img.size
                    dpi = img.info.get("dpi", (96, 96))
                    dpi_x = dpi[0] or 96
                    dpi_y = dpi[1] or 96

                ancho_in = ancho_px / dpi_x
                alto_in = alto_px / dpi_y

                max_width_in = max_width / 914400
                max_height_in = max_height / 914400

                escala = min(max_width_in / ancho_in, max_height_in / alto_in, 1.0)

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                if i > 0:
                    p.paragraph_format.page_break_before = True

                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(ancho_in * escala), height=Inches(alto_in * escala))

            if output_name:
                docx_path = output_dir / f"{output_name}.docx"
            else:
                docx_path = output_dir / "output.docx"

            doc.save(docx_path)
            output_files.append(docx_path)

            return OperationResult(True, f"Creado Word con {len(input_paths)} imagen(es)", output_files)
        except Exception as e:
            return OperationResult(False, f"Error: {str(e)}")
